"""
Pre-chunked document processor for template-based uploads

This is supposed be used by an aware admin who only adds 
core text and not fluff/boilerplate. No deduplication or 
complex structural inspection needed.

Full text extraction is used over streaming assuming that
pre-chunked uploads will be smaller, dense files.

Template Structure (Users fill this):
---
CHUNK: 1.1 Risk Assessment
CATEGORY: compliance
---
Content of risk assessment section goes here.
Multiple paragraphs allowed.

---
CHUNK: 1.2 Customer Verification
CATEGORY: compliance
---
More content here...
"""
import re
import logging
import asyncio
import pymupdf
from io import BytesIO
from pathlib import Path
from docx import Document
from app.config.setting import settings
from typing import List, Dict, AsyncGenerator, Optional, Any, Callable, Tuple

logger = logging.getLogger(__name__)


class PreChunkedProcessor:
    """
    Process documents that follow the pre-chunked template
    Works with both PDF and DOCX
    """
        
    def __init__(self):
        # simple pattern with flexible whitespace 
        # handling for validating template markers
        self.chunk_marker = re.compile(
            r'---\s*\n'                        # Opening --- with optional spaces, then newline
            r'\s*CHUNK:\s*(.+?)\s*\n'          # CHUNK: with flexible whitespace
            r'(?:\s*CATEGORY:\s*(.+?)\s*\n)?'  # Optional CATEGORY with flexible whitespace
            r'\s*---\s*(?:\n|$)',              # Closing --- with optional newline
            re.MULTILINE
        )


    async def validate_format(self, file_path: Path, filename: str) -> Tuple[bool, str]:
        """
        Validate pre-chunked format and extract full text in single pass

        Returns: 
            (is_valid, full_text)
        """
        try:
            # could add size check, however, let admins decide

            # extract full text - single pass
            ext = Path(filename).suffix.lower()
            full_text = await self._extract_full_text(file_path, ext)
            
            if not full_text:
                logger.warning("Pre-chunked file has insufficient content")
                return False, ""
            
            logger.info(
                f"Pre-chunked format validated: "
                f"file size: {len(full_text) / 1024:.1f}KB"
            )

            return True, full_text

        except ValueError as e:
            # Validation failed
            logger.warning(f"Pre-chunked validation failed: {e}")
            return False, "" 
        except Exception as e:
            logger.error(f"Pre-chunked validation error: {e}")
            return False, ""
    
    
    async def _extract_full_text(self, file_path: Path, ext: str) -> str:
        """
        Extract text with early validation (fail fast - PDF only)
        
        Raises:
            ValueError: If validation fails
        """
        if ext == '.pdf':
            return await self._extract_pdf(file_path)
        elif ext == '.docx':
            return await self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type for pre-chunked: {ext}")

    
    async def _extract_pdf(self, file_path: Path) -> str:
        """Extract PDF with early validation"""
        try:
            doc = pymupdf.open(str(file_path))
            total_pages = doc.page_count
            pages_text = []
            
            # Extract first 2 pages for validation
            validation_pages = min(2, total_pages)
            validation_text = ""
            
            for page_num in range(validation_pages):
                page_text = doc[page_num].get_text("text")
                pages_text.append(page_text)
                validation_text += page_text + "\n\n"
            
            # Early validation check
            if len(validation_text.strip()) < 100:
                doc.close()
                raise ValueError("Insufficient content in first pages")
            
            # Check for markers in first pages
            matches = list(self.chunk_marker.finditer(validation_text))
            
            logger.debug(
                f"Validation check: {len(matches)} markers found "
                f"in first {validation_pages} pages"
            )
            
            if len(matches) < settings.processing.MIN_MARKER_REQUIRED:
                doc.close()
                raise ValueError(
                    f"No valid chunk markers found in first {validation_pages} pages. "
                    f"Expected format:\n---\nCHUNK: Section Name\nCATEGORY: category\n---"
                )
            
            # Validation passed - extract remaining pages
            logger.info(
                f"Pre-chunked markers found, extracting remaining "
                f"{total_pages - validation_pages} pages"
            )
            
            for page_num in range(validation_pages, total_pages):
                pages_text.append(doc[page_num].get_text("text"))
            
            # cleanup
            doc.close()
            
            return "\n\n".join(pages_text)
            
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract PDF text: {e}")

    
    async def _extract_docx(self, file_path: Path) -> str:
        """
        Extract DOCX and validate format
        
        Note: DOCX paragraphs are already loaded into memory by python-docx,
        so no benefit to "early" validation - just extract all and validate.
        """
        try:
            doc = Document(str(file_path))
            
            # Extract all paragraphs at once (python-docx loads entire XML anyway)
            all_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            if not all_paragraphs:
                raise ValueError("No content found in document")
            
            full_text = "\n\n".join(all_paragraphs)
            
            # Validate after extraction
            if len(full_text.strip()) < 100:
                raise ValueError("Insufficient content in document")
            
            # Check for markers
            matches = list(self.chunk_marker.finditer(full_text))
            
            logger.debug(
                f"DOCX validation: {len(matches)} markers found "
                f"in {len(all_paragraphs)} paragraphs"
            )
            
            if len(matches) < settings.processing.MIN_MARKER_REQUIRED:
                # Show helpful error with sample
                sample = full_text[:500] if len(full_text) > 500 else full_text
                logger.debug(f"First 500 chars of document:\n{sample}")
                
                raise ValueError(
                    f"No valid chunk markers found in document. "
                    f"Expected format:\n---\nCHUNK: Section Name\nCATEGORY: category\n---"
                )
            
            logger.info(
                f"DOCX pre-chunked validated: {len(matches)} chunks, "
                f"{len(all_paragraphs)} paragraphs"
            )
            
            return full_text
            
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to extract DOCX text: {e}")


    async def process_file( # uses same name as main processor for unified caller
        self,
        text: str,
        upload_id: str,
        progress_callback: Optional[Callable[[int, int, str, str], None]] = None
    ) -> AsyncGenerator[tuple[str, bool], None]:
        """
        Process pre-chunked document with progress tracking
        
        Yields:
            (chunk_text, is_complete) tuples matching main processor interface
        """
        # Final validation check
        if not await self.is_prechunked(text):
            raise ValueError("Document does not follow pre-chunked template")
        
        # Split by chunk markers
        chunks = self._split_into_chunks(text)
        total_chunks = len(chunks)
    
        if total_chunks == 0:
            raise ValueError("No valid chunks found in document")

        logger.info(f"Processing {total_chunks} pre-defined chunks")
        
        for idx, chunk_data in enumerate(chunks):
            # Extract metadata
            section_id, section_title = self._parse_section_identifier(
                chunk_data['header']
            )
            
            # Format chunk with section header (matches standard chunker output)
            if section_id:
                formatted_chunk = f"{section_id}. {section_title}\n\n{chunk_data['content']}"
            else:
                formatted_chunk = f"{section_title}\n\n{chunk_data['content']}"
            
            # Progress callback
            if progress_callback:
                await progress_callback(
                    idx + 1,  # current
                    total_chunks,  # total
                    f"Processing {idx + 1} chunk",  # message
                    upload_id
                )
            
            # Yield
            is_complete = (idx == total_chunks - 1)
            yield formatted_chunk, is_complete
            
            # Yield control periodically
            if idx % 10 == 0:
                await asyncio.sleep(0.001)
    

    def _split_into_chunks(self, text: str) -> List[Dict[str, Any]]:
        """
        Split text into chunks based on markers
        """
        chunks = []
        
        # Find all chunk markers
        marker_positions = [
            (m.start(), m.group(1), m.group(2))
            for m in self.chunk_marker.finditer(text)
        ]
    
        if not marker_positions:
            logger.warning("No chunk markers found in text")
            return []

        # Extract chunks
        for idx, (start_pos, chunk_id, category) in enumerate(marker_positions):
            # Find content start (after marker)
            content_start = text.find('---', start_pos) + 3
            content_start = text.find('\n', content_start) + 1
            
            # Find content end (next marker or EOF)
            if idx + 1 < len(marker_positions):
                content_end = marker_positions[idx + 1][0]
            else:
                content_end = len(text)
            
            # Extract content
            content = text[content_start:content_end].strip()

            # Skip empty chunks
            if not content:
                logger.warning(f"Empty chunk found: {chunk_id}")
                continue

            chunks.append({
                'header': chunk_id.strip(),
                'category': category.strip() if category else None,
                'content': content
            })
        
        return chunks
    

    def _parse_section_identifier(self, header: str) -> tuple[str, str]:
        """
        Parse section header into ID and title
        
        Examples:
            "1.1 Risk Assessment" → ("1.1", "Risk Assessment")
            "4. CUSTOMER POLICY" → ("4", "CUSTOMER POLICY")
            "Introduction" → (None, "Introduction")
        """
        # Try numbered pattern
        match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', header)
        
        if match:
            return match.group(1), match.group(2).strip()
        
        # Fallback: entire header is title
        return None, header.strip()


    async def is_prechunked(self, text: str) -> bool:
        """
        Quick check if text follows pre-chunked template
        """
        matches = self.chunk_marker.findall(text)
        return len(matches) >= settings.processing.MIN_MARKER_REQUIRED
    

    def cleanup(self):
        """No resources to clean (interface compatibility)"""
        pass


def generate_template_docx() -> bytes:
    """
    Generate downloadable pre-chunk template (DOCX)
    
    Returns: Binary DOCX file
    """
    doc = Document()
    
    # Title
    doc.add_heading('Pre-Chunked Document Template', 0)
    
    # Instructions
    doc.add_paragraph(
        'Fill in your content below. Each section will become a separate searchable chunk.'
    )
    doc.add_paragraph(
        'Instructions:'
    )
    doc.add_paragraph(
        '• Use the markers exactly as shown (---)',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Section IDs should be numbered (1, 1.1, 2, etc.)',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Category is optional but recommended',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Each chunk should contain meaningful content (minimum 50 characters)',
        style='List Bullet'
    )
    
    doc.add_paragraph('')
    doc.add_heading('Example Chunks:', level=2)
    
    # Example 1
    doc.add_paragraph('---')
    doc.add_paragraph('CHUNK: 1. Introduction')
    doc.add_paragraph('CATEGORY: general')
    doc.add_paragraph('---')
    doc.add_paragraph(
        'Your introduction content goes here. Explain the purpose and scope of this document. '
        'You can use multiple paragraphs and formatting as needed.'
    )
    doc.add_paragraph('')
    
    # Example 2
    doc.add_paragraph('---')
    doc.add_paragraph('CHUNK: 2.1 Risk Assessment')
    doc.add_paragraph('CATEGORY: compliance')
    doc.add_paragraph('---')
    doc.add_paragraph(
        'Describe your risk assessment procedures. Include details about:\n'
        '• Risk identification criteria\n'
        '• Assessment methodology\n'
        '• Documentation requirements'
    )
    doc.add_paragraph('')
    
    # Example 3
    doc.add_paragraph('---')
    doc.add_paragraph('CHUNK: 2.2 Customer Verification')
    doc.add_paragraph('CATEGORY: compliance')
    doc.add_paragraph('---')
    doc.add_paragraph(
        'Your customer verification procedures go here. Be specific about required documents '
        'and verification steps.'
    )
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()