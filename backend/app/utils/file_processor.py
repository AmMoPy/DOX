import gc
import pymupdf
import logging
import asyncio
from uuid import UUID
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.document import CT_Body
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pathlib import Path
from typing import AsyncGenerator, Tuple, Optional, Callable
from fastapi import HTTPException
from app.utils.file_processor_utils import (
    BaseChunker, 
    BaseNormalizer, 
    BaseDeduplicator, 
    BaseHFdetector, 
    BaseStringCache,
    BaseFuzzyMatcher
    )
from app.utils.mem_mngr import AdaptiveMemoryManager
from app.db.db_factory import hash_store, doc_store
from app.config.setting import settings

logger = logging.getLogger(__name__)


class FileProcessor:
    """
    Main file upload processor that handles various document formats and structures

    Core Features:
    --------------
    1. Structural chunking - Automatically analyzes document structure (PDF/DOCX), 
    extracts content, detects sections/headers, and chunks based on semantic boundaries
    2. Streaming/Iterative - Processes documents page-by-page (PDF) or element-by-element 
    (DOCX) to minimize memory usage
    3. Sophisticated memory management with buffers, garbage collection, and adaptive flushing
    4. Complex with multiple components (chunker, normalizer, deduplicator, fuzzy matcher, 
    header/footer detection, string caching)

    Output:
    -------
    Chunks created dynamically based on content analysis and chunker settings
    """
    
    def __init__(self,
                 chunker: BaseChunker, # Using abstract instead of concrete as type hint because these are volatile dependencies having multiple implementations
                 normalizer: BaseNormalizer,
                 deduplicator: BaseDeduplicator,
                 fuzzy_matcher: BaseFuzzyMatcher,
                 mem_mngr: AdaptiveMemoryManager,
                 user_id: UUID,
                 document_id: str,
                 upload_id: str,
                 hfdetector: Optional[BaseHFdetector] = None,
                 scache: Optional[BaseStringCache] = None,
                 progress_callback: Optional[Callable[[int, int, str, str], None]] = None
                 ):
        # dependency injection pattern is better here as
        # each class has distinct responsibilities and
        # inheritance would create unnecessary coupling
        self.chunker = chunker
        self.normalizer = normalizer
        self.deduplicator = deduplicator
        self.fuzzy_matcher = fuzzy_matcher
        self.memory_manager = mem_mngr
        self.string_cache = scache
        self.hf_detector = hfdetector
        self.user_id = user_id
        self.document_id = document_id
        self.upload_id = upload_id
        self.progress_callback = progress_callback

        # Memory management settings
        # BUFFER LIMITS (in characters)
        # Smaller buffer for better section boundary detection
        self.max_buffer_size_dedup = 2_000       # Max chars in buffer (~2KB) to deduplicate, < process buffer for single tracking, cleaning and faster flow 
        self.max_buffer_size_process = 4_000     # Max chars in buffer (~4KB) to process
        self.max_buffer_size = 500_000           # 500KB text = ~100,000 words absolute max for edge cases of one giant chunk
        self.max_table_chars = 50_000            # 50KB text = ~10,000 words (assuming 5 chars per word) words_per_page = 500  # Typical book page
        self.max_cell_chars = 500                # Char limit per table cell (fine-grained)
        self.emergency_flush_threshold = 0.9     # Trigger approaching max_buffer_size warning 
        
        # COUNT LIMITS (element counts, not size)
        self.max_page_buffer = 2                   # Max pages to buffer before processing
        self.max_paragraph_buffer = 100            # Max paragraph for docx 

        # QUALITY FILTERS (in characters)
        self.skip_size = 50                        # Skip tiny text blocks

        # Processing state
        self.text_buffer = []
        self.buffer_size = 0
        self.document_fingerprint = None
        self.dedup_checked = False

        # header/footer state
        self.detected_headers = set()
        self.detected_footers = set()
        self.pages_for_detection = []  # Collect first N pages
    

    def _iter_docx_elements(self, doc: Document, parent_elm: CT_Body):
        """
        Iterate through all DOCX elements (paragraphs + tables) in document order
        """
        # Iterate through body elements in order
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)


    def _extract_element_text(self, element) -> str:
        """
        Extract text from any DOCX element (paragraph or table)
        """
        if hasattr(element, 'text'):  # Paragraph
            return element.text.strip()
        elif hasattr(element, 'rows'):  # Table
            return self._extract_table_text(element)
        return ""
    

    def _is_header_table(self, table) -> bool:
        """
        Quick check if table is likely a header/footer
        Good for inline headers disguised as tables
        """
        if len(table.rows) > 5:
            return False  # Headers are typically short
        
        # Check first row
        first_row = table.rows[0]
        all_short = all(
            len(''.join(para.text for para in cell.paragraphs)) < 100
            for cell in first_row.cells
        )
        
        return all_short and len(table.rows) <= 3


    def _extract_table_text(self, table, max_rows: int = 1000, max_cells_per_row: int = 50) -> str:
        """
        Extract structured text from tables with size limits and early exit 
        while preserving relationships, Prevents memory exhaustion from massive 
        tables

        Multi-level safety limits:
            1. Row count limit (prevents infinite loops)
            2. Cell per row limit (prevents wide tables)
            3. Total character limit (respects buffer processing limits)

        Returns:
            Extracted text + truncation notice if limits hit
        """
        # Quick check: Is this a header table? (all short text)
        if self._is_header_table(table):
            logger.debug("Skipping header table")
            return ""

        table_lines = []
        total_chars = 0
      
        # Process rows with hard limits
        for row_idx, row in enumerate(table.rows):
            # ROW LIMIT CHECK
            if row_idx >= max_rows:
                table_lines.append(f"... (table truncated at {max_rows} rows)")
                break # Breaks OUTER loop - stops ALL row processing

            row_data = []

            # Process cells with per-row limits
            for cell_idx, cell in enumerate(row.cells):
                # CELLS LIMIT (PER ROW) - graceful exit for current row
                if cell_idx >= max_cells_per_row:
                    # row_data.append("... (row truncated)")
                    break # Breaks INNER loop - continues to NEXT ROW

                # Extract cell text with size limit
                cell_text = ' '.join(
                    para.text.strip()[:self.max_cell_chars]  # Limit cell content
                    for para in cell.paragraphs
                    if para.text.strip()
                )

                if cell_text:
                    row_data.append(cell_text)

            # Build row text
            if not row_data:
                continue  # Skip empty rows

            row_text = " | ".join(row_data)
            row_size = len(row_text) # chars not bytes

            # Total character limit (CHECK BEFORE APPEND)
            if total_chars + row_size >= self.max_table_chars: # TODO: use processing buffer instead?
                # Stop BEFORE exceeding limit
                table_lines.append(f"... (table too large: stopped at {total_chars} chars)")
                break # Stops ALL row processing
            
            # No limits reached
            # Safe to append
            table_lines.append(row_text)
            total_chars += row_size
        
        if table_lines:
            return "\n\n".join(table_lines)

        return ""


    def _clear_processed_elements(self, parent_elm, current_index: int):
        """
        Clear processed elements from XML tree to free memory
        Only keep last 50 elements for safety
        """
        try:
            # Get all child elements
            children = list(parent_elm)
            
            # Keep only recent elements (last 50)
            if len(children) > 50:
                elements_to_clear = children[:max(0, current_index - 50)]
                
                for elem in elements_to_clear:
                    try:
                        elem.clear()  # Clear element content
                        # Note: Can't remove from parent during iteration
                    except Exception:
                        pass
        except Exception as e:
            logger.debug(f"Element cleanup failed: {e}")


    async def _fast_dedup_check(self) -> None:
        """
        Lightweight content deduplication that works with both SQL and Postgress
        focusing on deduplicating format variations (A.pdf vs A.docx) with exact content 
        but subtle text processing structural differences

        Multi-stage deduplication:
        Stage 1: Exact hash
        Stage 2: LSH candidates retrieval and similarity calculations
        """
        if self.dedup_checked:
            return

        try:
            # Use first chunk as representative sample
            text_block = "\n\n".join(self.text_buffer)

            # Normalize for consistent comparison
            normalized_text = self.normalizer.clean_text_structured(
                text_block, 
                settings.processing.TXT_NORM_DEDUPLICATION
            )

            # Why LSH? Fast, low storage, good for near-duplicates
            self.document_fingerprint = self.deduplicator.generate_fingerprint(normalized_text)

            if not self.document_fingerprint:
                raise HTTPException(409, "Could not generate document fingerprint")

            # STAGE 1: Semantic match (fastest)
            cache_key = self.deduplicator.get_fingerprint_hash(self.document_fingerprint)

            # Single indexed lookup - O(log n)
            existing_content = await hash_store.check_content_hash_exists(cache_key)

            # TODO: create cash entry of the duplicate for future lookup?

            if existing_content:
                # Early exit - duplicate found, nothing stored
                raise HTTPException(
                    409, 
                    f"Content already exists: {existing_content['filename']} " 
                    f"(user: {self.user_id})"
                )

            # STAGE 2: LSH candidate retrieval (indexed lookup!)
            if settings.processing.ENABLE_FUZZY_CACHE_MATCHING:
                candidates = await hash_store.find_lsh_candidates(
                    self.document_fingerprint,
                    max_candidates = settings.processing.DOC_LSH_MAX_CANDIDATES
                )
                
                if not candidates:
                    logger.info("No LSH candidates found - document is unique")
                else:
                    logger.info(f"Starting full similarity check on {len(candidates)} candidates")

                    # Find best match
                    best_match = await self.fuzzy_matcher.find_best_match(
                        self.document_fingerprint,
                        candidates,
                        threshold=settings.processing.DOC_SIMILARITY_THRESHOLD,
                        mode="doc"
                    )
                    
                    if best_match:
                        # Duplicate detected
                        raise HTTPException(
                            409,
                            f"Similar content exists: {best_match['filename']} "
                            f"(similarity: {best_match['similarity_score']:.0%})"
                        )
            
            # Not a duplicate - store with LSH
            await hash_store.store_content_hash(
                cache_key,
                self.document_id,
                self.document_fingerprint
            )
            
            logger.info(f"Content fingerprint stored with LSH: (doc: {self.document_id})")

            # prevent further checks
            self.dedup_checked = True

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Content deduplication check failed: {e}")
            raise
            # TODO: Continue - deduplication failure shouldn't break upload?


    async def _safe_buffer_append(self, text: str) -> AsyncGenerator[str, None]:
        """
        Safe buffer append with size-aware handling
        
        Strategy:
        1. Normal text (total buffer < processing limit): Buffer normally
        2. Large text (total buffer > memory limit): Flush buffer first, then add
        3. Gigantic text (text itself > memory limit): Process directly, bypass buffer
        """

        # chars based size - assuming only English language used
        text_size = len(text)

        try:
            # Text itself exceeds memory limit (rare but catastrophic)
            if text_size >= self.max_buffer_size:
                logger.critical(
                    f"Oversized text detected: {text_size:,} chars "
                    f"(limit: {self.max_buffer_size:,}). Processing directly."
                )

                # Process directly without buffering
                # Let chunker handle splitting into smaller chunks
                async for chunk in self._generate_chunk(text = text):
                    if chunk and len(chunk.strip()) > self.skip_size: # chars OK here, Chunks are small (< 1KB typically)
                        yield chunk
                
                return  # Don't add to buffer

            # Check if adding would exceed memory limits
            if self.buffer_size + text_size >= self.max_buffer_size: # TODO: use processing buffer instead?
                logger.debug(
                    f"Buffer flush: {self.buffer_size:,} + {text_size:,} "
                    f">= {self.max_buffer_size_process:,} chars"
                )

                if self.text_buffer:
                    # Flush existing buffer
                    async for chunk in self._generate_chunk(clear_buffer = True):
                        # TODO: This breaks the current chunker section continuation logic 
                        if chunk and len(chunk.strip()) > self.skip_size: # chars OK here, Chunks are small (< 1KB typically)
                            yield chunk

            # safe to append
            self.text_buffer.append(text)
            self.buffer_size += text_size
            
            # Warn if approaching limits
            if self.buffer_size >= self.max_buffer_size * self.emergency_flush_threshold:
                memory_pct = (self.buffer_size / self.max_buffer_size) * 100
                logger.warning(f"Buffer at {memory_pct:.1f}% of memory limit")
        
        except Exception as e:
            logger.error(f"Buffer append failed: {e}")
            raise


    async def _generate_chunk(
        self,         
        normalize: bool = True,
        clear_buffer: bool = False,
        text: Optional[str] = None
        ) -> AsyncGenerator[str, None]:
        """
        Universal text processing with consistent normalization
        
        Args:
            normalize: Apply text normalization (default: True)
            clear_buffer: Clear text_buffer after processing (default: False)
            text: Raw text to process (optional)
        
        Yields:
            Processed and chunked text
        """
        if clear_buffer and not self.text_buffer:
            return
        
        if not text:
            # Combine buffer content with page separators
            text_block = "\n\n".join(self.text_buffer)
        else:
            text_block = text

        if normalize:
            processed_text = self.normalizer.clean_text_structured(text_block, settings.processing.TXT_NORM_CHUNKING)
        else:
            processed_text = text_block
        
        # Use context-aware chunker
        async for chunk in self.chunker.chunk_text_streaming(processed_text):
            if chunk and len(chunk.strip()) > self.skip_size:
                yield chunk
        
        # Clear buffer for
        # Normal processing
        if clear_buffer:
            self.text_buffer.clear()
            self.buffer_size = 0


    async def _stream_extract_pdf(self, file_path: Path) -> AsyncGenerator[str, None]:
        """
        PDF extraction with memory efficiency and topic awareness
        """
        logger.debug(f"Starting PDF extraction: {file_path}")
        
        try:
            doc = pymupdf.open(str(file_path)) # Loads file metadata + index
            total_pages = doc.page_count
            logger.info(f"Processing {total_pages} pages")
            
            if self.hf_detector:
                # Collect sample pages for pattern detection
                for page_num in range(min(settings.processing.HF_PAGE_SAMPLES, total_pages)):
                    page = doc[page_num]
                    page_text = page.get_text("text")
                    self.pages_for_detection.append(page_text)
                    page = None
                
                # Detect header/footer patterns
                self.detected_headers, self.detected_footers = \
                    self.hf_detector.detect_patterns(self.pages_for_detection)

                # clear sample buffer
                self.pages_for_detection.clear()

            for page_num in range(total_pages):
                try:
                    # Load one page at a time
                    page = doc[page_num]

                    # Extracts text + positioning
                    page_text = page.get_text("text") 
                    
                    # Close page to free memory immediately
                    page = None

                    # Progress callback
                    if self.progress_callback:
                        await self.progress_callback(
                            page_num + 1, 
                            total_pages, 
                            f"Processing {page_num + 1} page",
                            self.upload_id
                        )
                    
                    # String interning = Storing only ONE copy of 
                    # identical strings in memory. Repeated strings
                    # are not memory objects rather than pointers to
                    # the interned (cached) string
                    if self.string_cache and len(page_text) < 500:
                        # Intern repeated strings (headers, common phrases)
                        page_text = self.string_cache.get_or_intern(page_text)
                        
                    if page_text and page_text.strip():
                        # Remove headers
                        if self.hf_detector:
                            page_text = self.hf_detector.remove_hf(
                                page_text,
                                self.detected_headers,
                                self.detected_footers
                            )

                        # Safe buffer append with emergency flush
                        # Prevents OOM errors, caps memory per buffer
                        # Against maliciously crafted documents or
                        # Edge cases (96 small + 1 huge page)
                        # TODO: where first text block flushed 
                        # before deduplication (no buffer accumulation) but
                        # should be the same in docx so hashing wont be affected
                        async for chunk in self._safe_buffer_append(page_text): 
                            yield chunk  # Only yields if direct processing triggered
          
                    # content deduplication  
                    should_dedup = (
                        (self.buffer_size >= self.max_buffer_size_dedup or # 
                        page_num >= total_pages - 1) and 
                        not self.dedup_checked # prevents further checks when buffer clears
                        )

                    if should_dedup:
                        await self._fast_dedup_check() # Let exception propagate

                    # If we reach here, not a duplicate - process normally
                    # Process buffer when it gets large or we're at the end
                    should_process = (
                        len(self.text_buffer) >= self.max_page_buffer or      # Page count
                        self.buffer_size >= self.max_buffer_size_process or   # Character count
                        page_num >= total_pages - 1                           # End of doc
                    )
                    
                    if should_process:
                        async for chunk in self._generate_chunk(clear_buffer = True):
                            yield chunk

                    # Adaptive garbage collection based on memory pressure
                    if self.memory_manager.should_gc(page_num):
                        gc_time = self.memory_manager.collect_with_timing()
                        if gc_time > 0.1:  # Log slow GC
                            logger.warning(f"Slow GC: {gc_time:.3f}s")
                    
                    # Log memory stats periodically
                    if page_num % 50 == 0:
                        stats = self.memory_manager.get_gc_stats()
                        logger.debug(f"GC stats: {stats}")
                
                except HTTPException:
                    raise  # Re-raise to terminate generator when duplicates
                except Exception as page_error:
                    logger.error(f"Error processing page {page_num + 1}: {page_error}")
                    continue
            
            # Process any remaining buffer
            if self.text_buffer:

                # Progress callback
                if self.progress_callback:
                    await self.progress_callback(
                        total_pages, 
                        total_pages, 
                        f"Processed {total_pages} page",
                        self.upload_id
                    )

                async for chunk in self._generate_chunk(clear_buffer = True):
                    yield chunk
            
            # log final stats 
            cache_stats = self.string_cache.get_stats() if self.string_cache else {}
            mem_stats = self.memory_manager.get_gc_stats()

            logger.info(
                f"PDF extraction completed: {total_pages} pages, "
                f"Intern stats: {cache_stats}, "
                f"Memory stats: {mem_stats}"
            )

            # Explicit document cleanup
            doc.close()
            doc = None
            gc.collect()
        
        except HTTPException:
            # No logging for expected exceptions
            raise  # Let it bubble up to the endpoint
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
    

    async def _stream_extract_docx(self, file_path: Path) -> AsyncGenerator[str, None]:
        """
        DOCX extraction with buffering and topic awareness
        """
        logger.debug(f"Starting DOCX extraction: {file_path}")
        
        try:
            # DOCX is a ZIP file with XML content
            doc = Document(str(file_path))     # Loads entire document (entire XML tree)     

            total_elements = len(doc.paragraphs) + len(doc.tables)  # Not exact but for progress
            element_count = 0
        
            # Normalize to page equivalents for frontend progress reporting
            # Average: ~25-30 elements per page in typical documents
            # If progress is too fast/slow, adjust the divisor:
            # // 20 = slower progress (more "pages", each update is smaller % of total)
            # // 30 = faster progress (fewer "pages")
            divisor = 30
            estimated_pages = max(1, total_elements // divisor)

            # Get body element reference once
            parent_elm = doc.element.body

            logger.info(f"Processing {total_elements} elements (~{estimated_pages} pages)")

            # Iterates over loaded tree, stream paragraphs only, excludes headers footers
            # inline headers are handled during normalization and table extraction
            for element in self._iter_docx_elements(doc, parent_elm):
                element_count += 1 # counts all even empty text for correct checks
                try:
                    is_table = hasattr(element, 'rows')
                    element_text = self._extract_element_text(element)

                    # CRITICAL: Clear element reference immediately
                    if hasattr(element, '_element'):
                        element._element = None  # Break circular reference
                    element = None  # Explicit deletion

                    if self.progress_callback:
                        # Convert element count to page equivalent
                        current_page_equivalent = element_count // divisor

                        await self.progress_callback(
                            current_page_equivalent,
                            estimated_pages,  # estimate
                            f"Processing {element_count} element",
                            self.upload_id
                        )

                    if self.string_cache:
                        # Intern repeated strings (headers, common phrases)
                        # helpfull for repeated text such as inline disguised
                        # headers since we are not using header/footer detection for docx
                        element_text = self.string_cache.get_or_intern(element_text)

                    if not element_text:
                        continue

                    # Extract text from elements
                    if element_text:
                        # Safe buffer append with emergency flush
                        async for chunk in self._safe_buffer_append(element_text):
                            yield chunk

                    # content deduplication 
                    should_dedup = (
                        (self.buffer_size >= self.max_buffer_size_dedup or
                        element_count >= total_elements - 1) and
                        not self.dedup_checked
                        )

                    if should_dedup:
                        await self._fast_dedup_check()
                    
                    # Process buffer when appropriate
                    should_process = (
                        len(self.text_buffer) >= self.max_paragraph_buffer or 
                        self.buffer_size >= self.max_buffer_size_process or
                        element_count >= total_elements - 1
                    )
                    
                    # if should_process and paragraph_buffer:
                    if should_process:
                        async for chunk in self._generate_chunk(clear_buffer = True):
                            if chunk and len(chunk.strip()) > self.skip_size:
                                yield chunk
                    
                    # Aggressive cleanup every N elements
                    if element_count % 100 == 0:
                        # Force clear processed elements from parent
                        self._clear_processed_elements(parent_elm, element_count)
                        await asyncio.sleep(0.001)

                    # Adaptive garbage collection
                    if self.memory_manager.should_gc(element_count):
                        gc_time = self.memory_manager.collect_with_timing()
                        if gc_time > 0.1:  # Log slow GC
                            logger.warning(f"Slow GC: {gc_time:.3f}s")
                    
                    # Log memory stats periodically
                    if element_count % 500 == 0:
                        mem_stats = self.memory_manager.get_gc_stats()
                        logger.debug(f"GC stats: {mem_stats}")
                
                except HTTPException:
                    raise
                except Exception as elem_error:
                    logger.error(f"Error processing element {element_count}: {elem_error}")
                    continue
            
            # Process final buffer         
            if self.text_buffer:

                if self.progress_callback:
                    await self.progress_callback(
                        estimated_pages,
                        estimated_pages,  # estimate
                        f"Processed {total_elements} element",
                        self.upload_id
                    )

                async for chunk in self._generate_chunk(clear_buffer = True):
                    yield chunk

            # log final stats 
            cache_stats = self.string_cache.get_stats() if self.string_cache else {}
            mem_stats = self.memory_manager.get_gc_stats()

            logger.info(
                f"DOCX extraction completed: {element_count} elements, "
                f"Intern stats: {cache_stats}, "
                f"Memory stats: {mem_stats}"
            )

            # Explicit document cleanup
            doc._part = None
            doc._element = None
            doc = None
            gc.collect()

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise


    async def process_file(self, file_path: Path, filename: str) -> AsyncGenerator[Tuple[str, bool], None]:
        """
        Memory-efficient and topic aware streaming processor with completion signaling
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Starting processing: {filename}")
        
        try:
            extension = Path(filename).suffix.lower() # using sanitized name
            
            # Choose appropriate extractor
            if extension == '.pdf':
                extractor = self._stream_extract_pdf(file_path) # <- Read from .tmp path
            elif extension == '.docx':
                extractor = self._stream_extract_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
            
            # Stream chunks with completion tracking
            chunk_count = 0
            last_chunk = None

            # Track chunks to detect completion
            # look-ahead logic For chunks [A, B, C]:
            # Get chunk A → store in last_chunk (don't yield yet)
            # Get chunk B → yield A with is_complete=False → store B in last_chunk
            # Get chunk C → yield B with is_complete=False → store C in last_chunk
            # No more chunks → yield C with is_complete=True
            # TODO: single-chunk files but there's a 1-chunk delay in streaming
            async for chunk in extractor:
                if chunk and len(chunk.strip()) > self.skip_size:
                    # If we had a previous chunk, yield it as not complete
                    if last_chunk is not None:
                        chunk_count += 1
                        yield last_chunk, False                        

                        # Report chunking progress
                        if self.progress_callback:
                            try:
                                await self.progress_callback(
                                    chunk_count, 
                                    -1,  # Total unknown during streaming
                                    f"Processing {chunk_count} chunk",
                                    self.upload_id
                                )
                            except Exception as e:
                                logger.warning(f"Progress callback error: {e}")

                        # Memory management
                        if chunk_count % 20 == 0:
                            await asyncio.sleep(0.001)
                            if chunk_count % 100 == 0:
                                gc.collect()
                    
                    # Store current chunk for next iteration
                    last_chunk = chunk
            
            # Yield the final chunk with completion signal
            if last_chunk is not None:
                chunk_count += 1
                yield last_chunk, True

            # Final progress update
            if self.progress_callback:
                try:
                    await self.progress_callback(
                        chunk_count,
                        chunk_count,  # Now we know total
                        f"Processed {chunk_count} chunk",
                        self.upload_id
                    )
                except Exception as e:
                    logger.warning(f"Progress callback error: {e}")
            
            if chunk_count == 0:
                raise ValueError("No processable content found in document")
            
            logger.info(f"Processing completed: {filename} ({chunk_count} chunks)")
       
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            raise
    

    def cleanup(self):
        """Clean up processor resources"""
        try:
            self.text_buffer.clear()
            self.buffer_size = 0
            self.detected_headers = None
            self.detected_footers = None
            self.progress_callback = None
            gc.collect()
            logger.debug("file processor cleaned up")
        except Exception as e:
            logger.error(f"Error during cleanup: {e}")


def create_streaming_processor(
    chunker: BaseChunker,
    normalizer: BaseNormalizer,
    deduplicator: BaseDeduplicator,
    fuzzy_matcher: BaseFuzzyMatcher,
    mem_mngr: AdaptiveMemoryManager,
    user_id: UUID,
    document_id: str,
    upload_id: str,
    hfdetector: Optional[BaseHFdetector] = None,
    scache: Optional[BaseStringCache] = None,
    progress_callback: Optional[Callable[[int, int, str, str], None]] = None,
) -> FileProcessor:
    """
    Create a new processor instance with optional progress tracking and configurable chunking
    
    Args:
        chunker: chunker class instance
        normalizer: text normalizer class instance
        deduplicator: content deduplicator class instance
        fuzzy_matcher:fuzzy matcher class instance
        mem_mngr: memory manager class instance
        user_id: uploading user ID
        document_id: uploaded document ID
        upload_id: upload request ID
        hfdetector: header and footer detector class instance
        scache: string cache class instance
        progress_callback: function to call with progress updates
    """
    return FileProcessor(
        chunker, normalizer, deduplicator, fuzzy_matcher, mem_mngr, 
        user_id, document_id, upload_id, hfdetector, scache, progress_callback
        )